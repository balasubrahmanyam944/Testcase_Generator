# main.py
import os
import re
import json
import csv
import time
import threading
import yaml
import requests
import openai
import PyPDF2
import io
import docx
from docx import Document
from flask import Flask, jsonify, request, render_template, send_file
from flask_cors import CORS
from flask_socketio import SocketIO
from werkzeug.utils import secure_filename
from openapi_spec_validator import validate_spec
import ollama

# --- OpenAPILoader --------------------------------------------------
class OpenAPILoader:
    """
    Loads an OpenAPI specification from a local file or URL.
    """
    def __init__(self, source: str):
        self.source = source
        self.spec = None

    def load(self) -> dict:
        if self.source.startswith("http"):
            response = requests.get(self.source)
            if response.status_code != 200:
                raise Exception("Failed to load OpenAPI spec from URL")
            try:
                self.spec = response.json()
            except json.JSONDecodeError:
                self.spec = yaml.safe_load(response.text)
        else:
            with open(self.source, "r") as file:
                try:
                    self.spec = json.load(file)
                except json.JSONDecodeError:
                    self.spec = yaml.safe_load(file)
        return self.spec

# --- OpenAPIValidator --------------------------------------------------
class OpenAPIValidator:
    """
    Validates the OpenAPI spec using an open-source validator.
    """
    def __init__(self, spec: dict):
        self.spec = spec

    def validate(self) -> bool:
        validate_spec(self.spec)
        return True

# --- OpenAPIExtractor --------------------------------------------------
class OpenAPIExtractor:
    """
    Extracts endpoints and authentication details from the OpenAPI spec.
    """
    def __init__(self, spec: dict):
        self.spec = spec

    def extract_endpoints(self) -> list:
        endpoints = []
        paths = self.spec.get("paths", {})
        for path, methods in paths.items():
            for method, details in methods.items():
                endpoint_info = {
                    "endpoint": path,
                    "method": method.upper(),
                    "summary": details.get("summary", ""),
                    "parameters": details.get("parameters", []),
                    "request_body": details.get("requestBody", {}),
                    "responses": details.get("responses", {})
                }
                endpoints.append(endpoint_info)
        return endpoints

    def extract_authentication(self) -> dict:
        components = self.spec.get("components", {})
        return components.get("securitySchemes", {})

# --- RuleBasedTestCaseGenerator --------------------------------------------------
class RuleBasedTestCaseGenerator:
    """Enhanced rule-based test case generator with proper initialization"""
    def __init__(self, endpoints, positive_count=3, negative_count=3, edge_count=3):
        self.endpoints = endpoints
        self.positive_count = positive_count
        self.negative_count = negative_count
        self.edge_count = edge_count

    def generate_test_cases(self) -> list:
        test_cases = []
        for endpoint in self.endpoints:
            # Existing test case generation logic
            base_case = {
                "endpoint": endpoint["endpoint"],
                "method": endpoint["method"],
                "request_body": endpoint.get("request_body", {})
            }

            # Positive cases
            for i in range(self.positive_count):
                test_cases.append({
                    "description": f"Positive {i+1}: Valid {endpoint['method']} {endpoint['endpoint']}",
                    "category": "Functional",
                    "steps": [
                        f"Send valid {endpoint['method']} request to {endpoint['endpoint']}",
                        "Verify successful response"
                    ],
                    "expected": "200 OK with valid response structure",
                    "priority": "High",
                    "test_type": "Positive"
                })

            # Negative cases
            for i in range(self.negative_count):
                test_cases.append({
                    "description": f"Negative {i+1}: Invalid {endpoint['method']} {endpoint['endpoint']}",
                    "category": "Validation",
                    "steps": [
                        f"Send invalid {endpoint['method']} request to {endpoint['endpoint']}",
                        "Verify error handling"
                    ],
                    "expected": "4xx error with proper message",
                    "priority": "Medium",
                    "test_type": "Negative"
                })

            # Edge cases
            for i in range(self.edge_count):
                test_cases.append({
                    "description": f"Edge {i+1}: Boundary {endpoint['method']} {endpoint['endpoint']}",
                    "category": "Edge Cases",
                    "steps": [
                        f"Send boundary value request to {endpoint['endpoint']}",
                        "Verify boundary handling"
                    ],
                    "expected": "Proper boundary condition handling",
                    "priority": "Medium",
                    "test_type": "Edge Case"
                })
        
        return test_cases

# --- LLMTestCaseGenerator --------------------------------------------------
class LLMTestCaseGenerator:
    """
    Uses OpenAI's API to generate richer, LLM-driven test cases.
    For each endpoint, the LLM is instructed to generate multiple test cases.
    It expects the LLM to return a JSON array of test cases.
    """
    def __init__(self, openai_api_key: str, model: str = "gpt-3.5-turbo", positive_count: int = 3, negative_count: int = 3, edge_count: int = 3):
        openai.api_key = "asvasvasv"
        self.model = model
        self.positive_count = positive_count
        self.negative_count = negative_count
        self.edge_count = edge_count

    def generate_test_cases(self, endpoint_details: dict) -> list:
        prompt = self._construct_prompt(endpoint_details)
        print("generating  data for ", endpoint_details, "\n")
        try:
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert API tester.Generate test scenarios as a JSON array. Please output raw JSON without any markdown formatting or extra text"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000,
            )
            generated_text = response.choices[0].message.content.strip()
            cleaned_text = re.sub(r"^```(?:json)?", "", generated_text)
            cleaned_text = re.sub(r"```$", "", cleaned_text).strip()
            print(cleaned_text)
            # Attempt to parse the generated text as a JSON array
            test_cases = json.loads(cleaned_text)
            if not isinstance(test_cases, list):
                test_cases = [test_cases]
        except Exception as e:
            print(f"Error parsing JSON for endpoint {endpoint_details.get('endpoint')}: {e}")
            # Print the raw generated_text if available
            try:
                print("Raw response:", cleaned_text)
            except Exception:
                print("No raw response available.")
            # Fallback: create a default test case with error details.
            test_cases = [{
                "scenario": "default",
                "endpoint": endpoint_details.get("endpoint", ""),
                "method": endpoint_details.get("method", ""),
                "description": f"LLM generated test case (parsing error): {str(e)}",
                "expected_status": 200,
                "request_body": endpoint_details.get("request_body", {})
            }]
        return test_cases

    def _construct_prompt(self, endpoint_details: dict) -> str:
        total = self.positive_count + self.negative_count + self.edge_count
        prompt = (
            "Generate a JSON array of test cases for the given API endpoint. "
            f"There should be {self.positive_count} positive test cases, {self.negative_count} negative test cases, and {self.edge_count} edge test cases (a total of {total}).\n"
            "Each test case should be a JSON object with the following keys: 'scenario', 'endpoint', 'method', 'description', 'expected_status', and optionally 'request_body'.\n\n"
        )
        prompt += f"Endpoint: {endpoint_details.get('endpoint')}\n"
        prompt += f"Method: {endpoint_details.get('method')}\n"
        if endpoint_details.get("parameters"):
            prompt += f"Parameters: {json.dumps(endpoint_details.get('parameters'), indent=2)}\n"
        if endpoint_details.get("request_body"):
            prompt += f"Request Body: {json.dumps(endpoint_details.get('request_body'), indent=2)}\n"
        if endpoint_details.get("responses"):
            prompt += f"Responses: {json.dumps(endpoint_details.get('responses'), indent=2)}\n"
        prompt += (
            "\nGenerate the test cases as described:\n"
            "- Positive: Use valid inputs and expect a successful response (e.g., status 200).\n"
            "- Negative: Use invalid or missing inputs and expect an error response (e.g., status 400).\n"
            "- Edge: Use boundary or unusual inputs and expect an appropriate response.\n"
            "Return the result as a valid JSON array."
        )
        return prompt

# --- EnhancedAPITestExecutor --------------------------------------------------
class EnhancedAPITestExecutor:
    """
    Executes API tests based on the generated test cases.
    Supports GET, POST, PUT, and DELETE methods.
    """
    def __init__(self, base_url: str, test_cases: list):
        self.base_url = base_url
        self.test_cases = test_cases

    def run_tests(self) -> list:
        results = []
        for test in self.test_cases:
            method = test["method"]
            url = f"{self.base_url}{test['endpoint']}"
            payload = test.get("request_body", {}) or {}
            try:
                if method == "GET":
                    response = requests.get(url)
                elif method == "POST":
                    response = requests.post(url, json=payload)
                elif method == "PUT":
                    response = requests.put(url, json=payload)
                elif method == "DELETE":
                    response = requests.delete(url)
                else:
                    results.append({
                        "endpoint": test["endpoint"],
                        "method": method,
                        "expected_status": test["expected_status"],
                        "actual_status": "N/A",
                        "scenario": test.get("scenario", "N/A"),
                        "result": "SKIPPED (unsupported HTTP method)"
                    })
                    continue
                status_code = response.status_code
                passed = (status_code == test["expected_status"])
                result = {
                    "endpoint": test["endpoint"],
                    "method": method,
                    "scenario": test.get("scenario", "N/A"),
                    "expected_status": test["expected_status"],
                    "actual_status": status_code,
                    "result": "PASS" if passed else "FAIL"
                }
            except Exception as e:
                result = {
                    "endpoint": test["endpoint"],
                    "method": method,
                    "scenario": test.get("scenario", "N/A"),
                    "expected_status": test["expected_status"],
                    "actual_status": None,
                    "result": f"ERROR: {str(e)}"
                }
            results.append(result)
        return results

# Replace the existing LLMTestCaseGenerator with this Ollama-compatible version
# In main.py - Updated OllamaTestCaseGenerator class
class OllamaTestCaseGenerator:
    def __init__(self, model="llama3"):
        """Initialize Ollama model"""
        self.model = model

    def clean_gdd_text(self, gdd_text):
        """Extracts all major sections dynamically from any GDD format."""
        gdd_text = gdd_text[:30000]  # Increased limit to capture more data

        # Identify all key GDD sections dynamically
        pattern = re.compile(
            r"\n\s*(?=(Gameplay Mechanics|Physics|AI|Controls|UI/UX|Multiplayer|Scoring|Win/Loss|Levels|Combat|Progression|Abilities|Interactions))",
            re.IGNORECASE
        )

        sections = pattern.split(gdd_text)
        extracted_sections = [sections[i] + sections[i + 1] for i in range(1, len(sections) - 1, 2)]

        return "\n\n".join(extracted_sections) if extracted_sections else gdd_text  # Fallback to full text

    def extract_json(self, response_text):
        """Extracts only the JSON portion from a raw response."""
        json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
        return json_match.group(0) if json_match else None

    def generate_test_cases(self, gdd_text):
        """Generate maximum number of structured test cases dynamically"""
        gdd_cleaned = self.clean_gdd_text(gdd_text)

        prompt = f"""
        Analyze this Game Design Document (GDD) and generate **the maximum number of unique, gameplay-specific test cases**.
        The test cases must:
        - Cover **every possible gameplay feature** (mechanics, physics, AI, UI, multiplayer, etc.).
        - Be **as detailed as possible**, covering edge cases, interactions, and expected outcomes.
        - Use a **structured format** and **NO LIMIT on test case count**.

        ---
        **Extracted GDD Content for Processing:**
        {gdd_cleaned}
        ---

        Return **ONLY JSON** with the **maximum possible test cases** (NO extra text, explanations, or stopwords):

        {{
            "test_cases": [
                {{
                    "feature": "Feature Name",
                    "tests": [
                        "Test - Detailed interaction and expected outcome.",
                        "Test - Edge case with specific conditions.",
                        "Test - High-stress scenario with unique failure condition."
                    ]
                }}
            ]
        }}
        """

        for attempt in range(3):  # Retry up to 3 times if response is invalid
            try:
                response = ollama.chat(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "You are a test case generator. Always return JSON."},
                        {"role": "user", "content": prompt}
                    ]
                )

                print(f"🔥 Attempt {attempt + 1}: OLLAMA RAW RESPONSE:", response)

                result_text = response.get("message", {}).get("content", "").strip()
                json_text = self.extract_json(result_text)

                if not json_text:
                    print("⚠️ Invalid response format. Retrying...")
                    time.sleep(2)
                    continue

                parsed_json = json.loads(json_text)

                if "test_cases" in parsed_json and parsed_json["test_cases"]:
                    return parsed_json  # Return extracted test cases

            except json.JSONDecodeError as e:
                print(f"❌ JSON Decode Error: {str(e)}. Retrying...")
                time.sleep(2)

            except Exception as e:
                print(f"❌ Error generating test cases: {str(e)}. Retrying...")
                time.sleep(2)

        print("❌ All retries failed. Returning empty test cases.")
        return {"test_cases": []}

# --- Utility: Export to CSV --------------------------------------------------
def export_to_csv(data: list, filename: str):
    """
    Exports a list of dictionaries to a CSV file.
    """
    if not data:
        print("No data to export.")
        return
    keys = data[0].keys()
    with open(filename, 'w', newline='', encoding='utf-8') as output_file:
        dict_writer = csv.DictWriter(output_file, fieldnames=keys)
        dict_writer.writeheader()
        dict_writer.writerows(data)
    print(f"Exported {len(data)} records to '{filename}'.")

# --- Utility: Get Base URL from Spec ------------------------------------------
def get_base_url(spec: dict) -> str:
    """
    Determines the base URL from the OpenAPI spec.
    Supports both OpenAPI 3 (servers) and Swagger 2 (host, schemes, basePath).
    """
    if "servers" in spec and spec["servers"]:
        return spec["servers"][0]["url"].rstrip("/")
    elif "host" in spec:
        schemes = spec.get("schemes", ["https"])
        scheme = "https" if "https" in schemes else schemes[0]
        host = spec["host"]
        basePath = spec.get("basePath", "")
        return f"{scheme}://{host}{basePath}".rstrip("/")
    else:
        raise Exception("No server information found in the spec.")

# --- Flask Server ------------------------------------------------------------
# Flask and SocketIO Initialization
app = Flask(__name__)
CORS(app, resources={r"/generate": {"origins": "http://localhost:3000"}})
socketio = SocketIO(app, async_mode='gevent')

# Directory Configuration
UPLOAD_FOLDER = "D:/in"
DOWNLOAD_FOLDER = "D:/out"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER

# Function to read CSV and extract descriptions
def read_csv_descriptions(filename):
    descriptions = []
    with open(filename, mode='r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            descriptions.append(row['description'])
    return descriptions

def extract_text_from_file(content, file_type):
    """Unified text extraction with proper function names"""
    try:
        if file_type == 'pdf':
            return extract_text_from_pdf(content)
        elif file_type in ['doc', 'docx']:
            return extract_text_from_docx(content)
        elif file_type in ['txt', 'md']:
            return content.decode('utf-8')
        else:
            return str(content)
    except Exception as e:
        print(f"Text extraction failed: {str(e)}")
        return None

# Add missing PDF extraction function
def extract_text_from_pdf(content):
    """Extract text from PDF bytes"""
    try:
        with io.BytesIO(content) as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            text = "\n".join([page.extract_text() for page in reader.pages])
            return text
    except Exception as e:
        print(f"PDF extraction error: {str(e)}")
        return None

def extract_text_from_docx(content):
    """Extract text from DOCX bytes"""
    try:
        with io.BytesIO(content) as doc_file:
            doc = Document(doc_file)
            return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        return None

def extract_text_from_file(content, file_type):
    """Unified text extraction with proper error handling"""
    try:
        if file_type == 'pdf':
            return extract_text_from_pdf(content)
        elif file_type in ['doc', 'docx']:
            return extract_text_from_docx(content)
        elif file_type in ['txt', 'md', 'json']:
            return content.decode('utf-8')
        else:
            print(f"Unsupported file type: {file_type}")
            return None
    except Exception as e:
        print(f"Text extraction failed: {str(e)}")
        return None
       
def extract_text_from_document(file_path, file_type):
    """Extract text from GDD file based on its format."""
    try:
        if file_type == "txt":
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        elif file_type == "json":
            with open(file_path, "r", encoding="utf-8") as file:
                return json.dumps(json.load(file), indent=2)
        elif file_type == "docx":
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_type == "pdf":
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        else:
            return "Unsupported file format"
    except Exception as e:
        return f"Error reading document: {str(e)}"

def generate_test_cases_with_ollama(spec_data):
    """
    Generate test cases in original format
    """
    test_cases = {
        "llm_descriptions": [],
        "rule_based_descriptions": []
    }

    paths = spec_data.get("paths", {})
    for path, methods in paths.items():
        for method, details in methods.items():
            method = method.upper()
            
            # LLM-generated cases
            test_cases["llm_descriptions"].extend([
                f"Verify {method} {path} returns correct response format",
                f"Verify {method} {path} handles invalid authentication",
                f"Verify {method} {path} handles rate limiting"
            ])

            # Rule-based cases
            test_cases["rule_based_descriptions"].extend([
                f"Positive: Verify {method} {path} with valid parameters returns 200",
                f"Negative: Verify {method} {path} with missing parameters returns 400",
                f"Edge Case: Verify {method} {path} with maximum payload size"
            ])

    return test_cases

def generate_test_cases_from_document(document_text, file_type):
    """Generate dynamic test cases from any Game Design Document (GDD)."""
    try:
        prompt = f"""Analyze the following Game Design Document (GDD) and generate comprehensive test cases.
        Extract game-specific mechanics, interactions, and conditions, ensuring coverage of:
        - Core gameplay mechanics
        - Physics interactions
        - Controls and UI
        - Level objectives
        - AI behavior and pathfinding
        - Multiplayer/networking elements (if applicable)
        - Edge cases and failure scenarios

        GDD Content:
        {document_text[:20000]}

        Generate test cases in structured JSON format:
        {{
            "test_cases": [
                {{
                    "feature": "Feature Name",
                    "tests": [
                        "Test Case 1",
                        "Test Case 2",
                        "Test Case 3"
                    ]
                }}
            ]
        }}

        Only return valid JSON, with no extra text.
        """

        ollama = OllamaTestCaseGenerator()
        raw_response = ollama.generate_test_cases(prompt)

        # Debug: Print raw response from the AI
        print("RAW RESPONSE:", raw_response)

        # Extract test cases
        test_cases = raw_response.get("test_cases", [])

        # Flatten test cases
        llm_cases = [test for feature_group in test_cases for test in feature_group["tests"]]

        return {
            "llm_descriptions": [str(test) for test in llm_cases[:100]]
        }

    except Exception as e:
        print(f"Game test generation error: {str(e)}")
        return {
            "llm_descriptions": [f"Test generation error: {str(e)}"]
        }


@app.route('/generate', methods=['POST'])
def generate():
    try:
        # Handle JSON requests with spec_url
        if request.is_json:
            data = request.get_json()
            if 'spec_url' in data:
                loader = OpenAPILoader(data['spec_url'])
                spec = loader.load()
                return jsonify(generate_test_cases_with_ollama(spec))

        # Handle file uploads
        if 'file' in request.files:
            file = request.files['file']
            if file.filename == '':
                return jsonify({"error": "No selected file"}), 400
                
            file_type = file.filename.split('.')[-1].lower()
            content = file.read()  # Get binary content
            
            text = extract_text_from_file(content, file_type)
            
            if not text:
                return jsonify({"error": "Failed to extract text"}), 400
                
            return jsonify(generate_test_cases_from_document(text, file_type))

        return jsonify({"error": "Invalid request format"}), 400

    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({
            "error": f"Server error: {str(e)}",
            "llm_descriptions": [],
            "rule_based_descriptions": []
        }), 500

# File Upload API
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    if filename.endswith('.json'):
        processed_filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
        with open(filepath, "r") as f:
            file_data = f.read()
        with open(processed_filepath, "w") as f:
            f.write(file_data)
        socketio.emit('file_uploaded', {'filename': filename})
    return '', 204

# File Download API
@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404
    return send_file(filepath, as_attachment=True)

# File Delete API
@app.route('/delete/<filename>', methods=['DELETE'])
def delete_file(filename):
    filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
    if os.path.exists(filepath):
        os.remove(filepath)
        socketio.emit('file_deleted', {'filename': filename})
        return jsonify({"success": True})
    return jsonify({"error": "File not found"}), 404

# Monitor Download Folder
def monitor_download_folder():
    previous_files = set(os.listdir(DOWNLOAD_FOLDER))
    while True:
        time.sleep(1)
        current_files = set(os.listdir(DOWNLOAD_FOLDER))
        deleted_files = previous_files - current_files
        for file in deleted_files:
            socketio.emit('file_deleted', {'filename': file})
        new_files = current_files - previous_files
        for file in new_files:
            if file.endswith('.json'):
                socketio.emit('file_uploaded', {'filename': file})
        previous_files = current_files

if __name__ == '__main__':
    thread = threading.Thread(target=monitor_download_folder)
    thread.daemon = True
    thread.start()
    socketio.run(app, debug=True, host='0.0.0.0', port=8081)


# URL: https://petstore.swagger.io/v2/swagger.json